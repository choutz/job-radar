import io
import os
import streamlit as st
from datetime import date
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from dotenv import load_dotenv

from models import get_config, Job

load_dotenv()


def _get_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY")
    return Anthropic(api_key=api_key)


def _set_font(run, bold=False, size_pt=11):
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Also set the theme font so Word doesn't override it
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def _add_tight_paragraph(doc, text, bold=False, size_pt=11):
    """Add a paragraph with no extra before/after spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_font(run, bold=bold, size_pt=size_pt)
    return p


def _add_body_paragraph(doc, text, size_pt=11):
    """Add a body paragraph with standard spacing after."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt)
    return p


def generate_cover_letter_text(job: Job) -> str:
    user_profile = get_config("USER_PROFILE")
    prompt = f"""
{user_profile}

Write a cover letter body for the following job posting.
Write only the body paragraphs (3-4 paragraphs).
Do not include a date, greeting, or signature block — just the body text, with paragraphs separated by a single blank line.
Be specific to this role and company. Be concise and professional.

Job Title: {job.title}
Company: {job.company}
Job Description: {job.description[:10000] if job.description else 'N/A'}
"""
    response = _get_client().messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}],
        system="You are a professional cover letter writer. Write compelling, tailored cover letters.",
    )
    return response.content[0].text.strip()


def build_cover_letter_doc(job: Job) -> bytes:
    name = get_config("COVER_LETTER_NAME")
    contact = get_config("COVER_LETTER_CONTACT")
    body_text = generate_cover_letter_text(job)

    doc = Document()

    # Margins (match PDF: ~1" top/bottom, ~1.25" sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # --- Header block (tight, no spacing between lines) ---
    _add_tight_paragraph(doc, name)
    _add_tight_paragraph(doc, contact)
    _add_tight_paragraph(doc, date.today().strftime("%B %d, %Y"))
    _add_tight_paragraph(doc, job.company)

    # Blank line after header
    _add_tight_paragraph(doc, "")

    # Greeting
    _add_tight_paragraph(doc, "Dear Hiring Manager,")

    # Blank line before body
    _add_tight_paragraph(doc, "")

    # --- Body paragraphs ---
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    for i, para in enumerate(paragraphs):
        space_after = Pt(10) if i < len(paragraphs) - 1 else Pt(0)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = space_after
        run = p.add_run(para)
        _set_font(run)

    # Blank line before closing
    _add_tight_paragraph(doc, "")

    # Closing
    _add_tight_paragraph(doc, "Sincerely,")
    _add_tight_paragraph(doc, name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()